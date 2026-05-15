import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_code_block(paragraph, code):
    """
    Add code to a Word paragraph using monospaced font.
    """
    run = paragraph.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    # Ensure correct font rendering in Word
    rPr = run._element.rPr
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Courier New')
    rPr.append(rFonts)

def write_source_files_to_docx(base_folder, output_filename='source_code.docx'):
    """
    Recursively find all .cpp, .c, and .hpp files in base_folder,
    and write their contents into a Word (.docx) file.
    """
    document = Document()
    extensions = ('.cpp', '.c', '.hpp')

    for root, dirs, files in os.walk(base_folder):
        for file in sorted(files):
            if file.endswith(extensions):
                full_path = os.path.join(root, file)

                # Get folder path without filename
                rel_dir = os.path.relpath(root, base_folder)
                location = base_folder if rel_dir == '.' else os.path.join(base_folder, rel_dir)

                # Add filename as heading
                document.add_heading(file, level=1)

                # Add folder path (not including filename)
                document.add_paragraph(f'Location: {location}', style='Intense Quote')

                # Read and insert code content
                try:
                    with open(full_path, 'r', encoding='utf-8', errors='ignore') as f:
                        code = f.read()
                except Exception as e:
                    document.add_paragraph(f"Error reading file: {e}")
                    continue

                code_paragraph = document.add_paragraph()
                add_code_block(code_paragraph, code)

    # Save the final document
    document.save(output_filename)
    print(f"\n✅ Source code exported to: {output_filename}")

# ========== USAGE ==========

# Replace this path with the root folder of your source files
source_folder = '/home/abhijith/Desktop/dev/bookin/booksim2/src'  # Or use absolute path
output_file = 'source_code.docx'

write_source_files_to_docx(source_folder, output_file)

